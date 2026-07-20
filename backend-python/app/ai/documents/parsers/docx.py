"""python-docx-backed DOCX parser."""

from __future__ import annotations

import asyncio
import io

from docx import Document as DocxDocument

from app.ai.documents.schemas import ParsedDocument


class DocxParser:
    async def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(self._parse_sync, file_bytes, filename)

    def _parse_sync(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        document = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [
            paragraph.text for paragraph in document.paragraphs if paragraph.text
        ]
        text = "\n".join(paragraphs)
        return ParsedDocument(
            text=text,
            metadata={
                "source": filename,
                "paragraph_count": len(paragraphs),
            },
        )
